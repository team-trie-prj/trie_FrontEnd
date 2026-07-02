# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_korean_font(doc, font_name='맑은 고딕'):
    # Set default font for normal text
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_report():
    doc = docx.Document()
    set_korean_font(doc)
    
    # Title
    title = doc.add_heading('프로젝트 진행 현황 및 잔여 태스크 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set font for heading
    for run in title.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    doc.add_paragraph('프로젝트명: 생성형 AI 기반 지능형 정보 시스템 설계 및 연구').bold = True
    doc.add_paragraph('작성일: 2026년 6월 21일')
    doc.add_paragraph('')
    
    # Current Status
    h1 = doc.add_heading('1. 현재 진행 상태 (Current Status)', level=1)
    for run in h1.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    p = doc.add_paragraph()
    p.add_run('현재 프로젝트는 핵심 모듈에 대한 프로토타입 검증 단계가 일부 완료된 상태입니다.').bold = True
    
    doc.add_paragraph('하이브리드 RAG 아키텍처 프로토타입 구현: 의미 검색(임베딩)과 키워드 검색(BM25)을 결합하고 RRF로 재정렬하는 파이프라인 개발 완료.', style='List Bullet')
    doc.add_paragraph('통합 스키마(UnifiedDoc) 설계: 사내 비전 데이터와 공공데이터포털 연계 데이터를 통합 관리하기 위한 SQLite 기반 데이터베이스 스키마 및 적재 로직 구현.', style='List Bullet')
    doc.add_paragraph('AI 에이전트 기반 구조 설계: 질의 의도 분석, 도구 호출(Function-Calling), 오프라인 규칙 라우터를 수행하는 Gemini 기반 LLM 에이전트 루프 구축 완료.', style='List Bullet')
    doc.add_paragraph('기능 데모(PoC) 확보: 대전 공공데이터(유성구 포트홀 등) 샘플 데이터를 통한 질의응답 및 검색 증강 시나리오 동작 확인.', style='List Bullet')
    
    # Remaining Tasks
    h2 = doc.add_heading('2. 잔여 태스크 (Remaining Tasks)', level=1)
    for run in h2.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    p2 = doc.add_paragraph()
    p2.add_run('제안서 최종본 및 현재 요구사항을 바탕으로 완성도 있는 프로토타입/보고서 산출을 위해 다음 작업들이 필요합니다.').bold = True
    
    doc.add_paragraph('실제 공공데이터/비전 데이터 파이프라인(ETL) 연동: 샘플 JSONL을 넘어 실제 CSV 등 원본 데이터를 정규화하고, 좌표 기반 공간 조인 등 데이터 처리 로직 완성.', style='List Bullet')
    doc.add_paragraph('VLM 적용 및 이미지 라벨링 자동화: Vision AI 솔루션 데이터에 특화된 비전언어모델(VLM)을 활용한 이미지 자동 이해 및 분석 파이프라인 개발.', style='List Bullet')
    doc.add_paragraph('문서/보고서 자동 생성 모듈 개발: 검색 결과와 데이터 기반의 통계를 요약하여 최종적으로 보고서 템플릿(.docx 형태 등)으로 자동 산출하는 기능 구현.', style='List Bullet')
    doc.add_paragraph('UI/UX 설계안 구성 및 화면 프로토타입 구현: 자연어 질의 화면, 검색 결과 화면, 데이터 요약 화면, 공공데이터 연계 화면 등 서비스 흐름에 맞춘 화면 UI 구성.', style='List Bullet')
    doc.add_paragraph('데이터/이력 관리 체계 구축: 데이터 변경 이력, 프롬프트, 응답 결과, 수정 이력 등을 관리할 수 있는 구조 및 화면 기획.', style='List Bullet')
    doc.add_paragraph('문서화 산출물 완성: 기능 정의서, 통합형 AI 시스템 아키텍처 설계안, 향후 확장 방향 등을 정리한 최종 기획 보고서 및 발표자료(PPT) 작성.', style='List Bullet')
    doc.add_paragraph('파인튜닝 및 모듈 종합 검증: 도메인 특화 성능 향상을 위한 파인튜닝 검토 및 전체 에이전트 파이프라인 테스트 시나리오(업무 자동화 추천 등) 검증.', style='List Bullet')

    doc.add_paragraph('\n')
    doc.add_paragraph('본 보고서는 코드 저장소(src)의 진행 내역과 제안서(PDF)의 기획 범위를 비교 분석하여 작성되었습니다.', style='Intense Quote')
    
    # Save document
    doc.save('Project_Remaining_Tasks_Report.docx')
    print("Report saved to Project_Remaining_Tasks_Report.docx")

if __name__ == "__main__":
    create_report()
